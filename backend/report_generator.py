from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn, nsdecls
from docx.oxml import OxmlElement, parse_xml
import json
from datetime import datetime
import os
import graphviz

class VAPTReportGenerator:
    def __init__(self, report_data):
        self.report_data = report_data
        self.document = Document()
        self.setup_document()

    def setup_document(self):
        # Set up document properties
        self.document.core_properties.author = self.report_data["metadata"]["author"]
        self.document.core_properties.title = self.report_data["metadata"]["title"]
        self.document.core_properties.subject = self.report_data["metadata"]["subject"]
        
        # Set up sections
        section = self.document.sections[0]
        section.page_height = Inches(11)
        section.page_width = Inches(8.5)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        
        # Set up styles
        styles = self.document.styles
        
        # Heading styles
        for i in range(1, 5):
            style_name = f'Heading {i}'
            if style_name in styles:
                style = styles[style_name]
                style.font.size = Pt(18 - (i * 2))  # Decreasing size for deeper headings
                style.font.bold = True
                if i == 1:
                    style.font.color.rgb = RGBColor(0, 51, 102)  # Dark blue for main headings
        
        # Normal text style
        if 'Normal' in styles:
            style = styles['Normal']
            style.font.name = 'Calibri'
            style.font.size = Pt(11)
            
        # List Bullet style
        if 'List Bullet' in styles:
            style = styles['List Bullet']
            style.font.name = 'Calibri'
            style.font.size = Pt(11)
        
        # Create a custom style for code blocks
        try:
            code_style = styles.add_style('CodeBlock', WD_STYLE_TYPE.PARAGRAPH)
            code_style.font.name = 'Courier New'
            code_style.font.size = Pt(9)
            code_style.paragraph_format.space_before = Pt(6)
            code_style.paragraph_format.space_after = Pt(6)
            code_style.paragraph_format.left_indent = Pt(12)
            
            # Add light gray shading to code blocks
            element = parse_xml(f'<w:shd {nsdecls("w")} w:fill="E6E6E6"/>')
            code_style.element.rPr.append(element)
        except:
            # Style might already exist
            if 'CodeBlock' in styles:
                code_style = styles['CodeBlock']
                code_style.font.name = 'Courier New'
                code_style.font.size = Pt(9)
                code_style.paragraph_format.space_before = Pt(6)
                code_style.paragraph_format.space_after = Pt(6)
                code_style.paragraph_format.left_indent = Pt(12)
                
                # Add light gray shading to code blocks
                element = parse_xml(f'<w:shd {nsdecls("w")} w:fill="E6E6E6"/>')
                code_style.element.rPr.append(element)

    def add_title_page(self):
        # Add company logo placeholder
        # self.document.add_picture('logo.png', width=Inches(2))
        
        # Add title
        title = self.document.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title.add_run(self.report_data["content"]["title"])
        title_run.font.size = Pt(24)
        title_run.font.bold = True
        
        # Add subtitle
        subtitle = self.document.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle_run = subtitle.add_run("Security Assessment Report")
        subtitle_run.font.size = Pt(18)
        
        # Add metadata
        metadata = self.document.add_paragraph()
        metadata.alignment = WD_ALIGN_PARAGRAPH.CENTER
        metadata_text = [
            f"Prepared by: {self.report_data['metadata']['author']}",
            f"Assessment Date: {self.report_data['metadata']['assessment_date']}",
            f"Test Type: {self.report_data['metadata']['test_type']}"
        ]
        metadata.add_run("\n".join(metadata_text)).font.size = Pt(12)
        
        self.document.add_page_break()

    def add_executive_summary(self):
        self.add_heading("Executive Summary", 1)
        summary = self.document.add_paragraph()
        summary.add_run(self.report_data["content"]["executive_summary"])
        
        # Add findings statistics
        findings = self.report_data["content"]["findings"]
        stats = {
            "High": len([f for f in findings if f["severity"] == "High"]),
            "Medium": len([f for f in findings if f["severity"] == "Medium"]),
            "Low": len([f for f in findings if f["severity"] == "Low"])
        }
        
        self.document.add_paragraph()
        self.add_heading("Findings Overview", 2)
        table = self.document.add_table(rows=1, cols=3)
        table.style = "Table Grid"
        header_cells = table.rows[0].cells
        header_cells[0].text = "Severity"
        header_cells[1].text = "Count"
        header_cells[2].text = "Risk Level"
        
        for severity, count in stats.items():
            row_cells = table.add_row().cells
            row_cells[0].text = severity
            row_cells[1].text = str(count)
            row_cells[2].text = severity

    def add_findings(self):
        self.add_heading("Detailed Findings", 1)
        
        for idx, finding in enumerate(self.report_data["content"]["findings"], 1):
            # Finding header with severity indicator
            finding_title = f"{idx}. {finding['title']}"
            heading = self.add_heading(finding_title, 2)
            
            # Add severity with appropriate color
            severity_para = self.document.add_paragraph()
            severity_run = severity_para.add_run(f"Severity: {finding['severity']}")
            severity_run.bold = True
            
            # Set color based on severity
            if finding['severity'].lower() == 'high':
                severity_run.font.color.rgb = RGBColor(255, 0, 0)  # Red
            elif finding['severity'].lower() == 'medium':
                severity_run.font.color.rgb = RGBColor(255, 165, 0)  # Orange
            else:
                severity_run.font.color.rgb = RGBColor(0, 128, 0)  # Green
            
            # Summary Tab Content
            self.add_heading("Summary", 3)
            
            # Add description in a styled box
            desc_para = self.document.add_paragraph()
            desc_para.add_run("Vulnerability Description").bold = True
            self.document.add_paragraph(finding["description"])
            
            # Add affected URL if available
            if finding.get("affected_url"):
                url_para = self.document.add_paragraph()
                url_para.add_run("Affected URL: ").bold = True
                url_para.add_run(finding["affected_url"])
            
            # Replication Steps Tab Content
            if finding.get("replication_steps"):
                self.add_heading("Replication Steps", 3)
                steps = finding["replication_steps"]
                
                for step_num, step in enumerate(steps, 1):
                    # Add step number and title with better formatting
                    step_header = self.document.add_paragraph()
                    step_header.add_run(f"Step {step_num}: {step.get('title', '')}").bold = True
                    
                    # Add step description if available
                    if step.get('description'):
                        desc_para = self.document.add_paragraph()
                        desc_para.add_run(step['description'])
                    
                    # Add command in a code cell if available
                    if step.get('command'):
                        cmd_para = self.document.add_paragraph(style='CodeBlock')
                        cmd_para.add_run("Command:").bold = True
                        cmd_para.add_run("\n" + step['command'])
                    
                    # Add response in a code cell if available
                    if step.get('response'):
                        resp_para = self.document.add_paragraph(style='CodeBlock')
                        resp_para.add_run("Response:").bold = True
                        resp_para.add_run("\n" + step['response'])
                    
                    # Add findings/observations if available
                    if step.get('findings'):
                        find_para = self.document.add_paragraph()
                        find_para.add_run("Observations:").bold = True
                        
                        # Add each finding as a bullet point
                        for observation in step['findings']:
                            bullet_para = self.document.add_paragraph(style='List Bullet')
                            bullet_para.add_run(observation)
                    
                    # Add spacing between steps
                    self.document.add_paragraph()
            elif finding.get("proof_of_concept"):  # Fallback for older data format
                self.add_heading("Proof of Concept", 3)
                self.document.add_paragraph(finding["proof_of_concept"])
            
            # Risk & Impact Tab Content
            self.add_heading("Risk & Impact Analysis", 3)
            
            # CVSS Scores if available
            if finding.get("cvss_score"):
                self.add_heading("CVSS Scores", 4)
                cvss_table = self.document.add_table(rows=1, cols=3)
                cvss_table.style = "Table Grid"
                
                # Add header row
                header_cells = cvss_table.rows[0].cells
                header_cells[0].text = "Base Score"
                header_cells[1].text = "Temporal Score"
                header_cells[2].text = "Environmental Score"
                
                # Add scores row
                score_row = cvss_table.add_row().cells
                score_row[0].text = str(finding["cvss_score"].get("base", "N/A"))
                score_row[1].text = str(finding["cvss_score"].get("temporal", "N/A"))
                score_row[2].text = str(finding["cvss_score"].get("environmental", "N/A"))
            
            # Business Impact
            impact_para = self.document.add_paragraph()
            impact_para.add_run("Business Impact:").bold = True
            self.document.add_paragraph(finding.get("impact", "No impact information available."))
            
            # Remediation
            self.add_heading("Remediation", 4)
            self.document.add_paragraph(finding.get("remediation", "No remediation steps available."))
            
            # Add reference if available
            if finding.get("reference"):
                ref_para = self.document.add_paragraph()
                ref_para.add_run("Reference: ").bold = True
                ref_para.add_run(finding["reference"])
            
            # Add MITRE ATT&CK mapping
            self.add_mitre_attack_mapping(finding)
            
            # Add developer recommendations
            self.add_developer_recommendations(finding)
            
            # Add page break between findings
            self.document.add_page_break()

    def add_mitre_attack_mapping(self, finding):
        """Add MITRE ATT&CK framework mapping for the finding"""
        self.add_heading("MITRE ATT&CK Mapping", 3)
        
        # Define attack stages based on finding type and severity
        attack_stages = []
        
        # Map finding types to MITRE tactics
        if "unauthenticated access" in finding["title"].lower() or "signalr" in finding["title"].lower():
            attack_stages = [
                {"tactic": "Initial Access", "technique": "Exploit Public-Facing Application (T1190)"},
                {"tactic": "Persistence", "technique": "External Remote Services (T1133)"},
                {"tactic": "Defense Evasion", "technique": "Exploitation for Privilege Escalation (T1068)"}
            ]
        elif "jquery" in finding["title"].lower() or "javascript" in finding["title"].lower():
            attack_stages = [
                {"tactic": "Initial Access", "technique": "Drive-by Compromise (T1189)"},
                {"tactic": "Execution", "technique": "Client-Side Execution (T1059.007)"},
                {"tactic": "Defense Evasion", "technique": "Obfuscated Files or Information (T1027)"}
            ]
        elif "information disclosure" in finding["title"].lower() or "server information" in finding["title"].lower():
            attack_stages = [
                {"tactic": "Reconnaissance", "technique": "Active Scanning (T1595)"},
                {"tactic": "Discovery", "technique": "System Information Discovery (T1082)"},
                {"tactic": "Collection", "technique": "Automated Collection (T1119)"}
            ]
        else:
            # Generic mapping for other findings
            if finding["severity"].lower() == "high":
                attack_stages = [
                    {"tactic": "Initial Access", "technique": "Valid Accounts (T1078)"},
                    {"tactic": "Execution", "technique": "Command and Scripting Interpreter (T1059)"},
                    {"tactic": "Persistence", "technique": "Create Account (T1136)"},
                    {"tactic": "Privilege Escalation", "technique": "Exploitation for Privilege Escalation (T1068)"}
                ]
            elif finding["severity"].lower() == "medium":
                attack_stages = [
                    {"tactic": "Initial Access", "technique": "Phishing (T1566)"},
                    {"tactic": "Execution", "technique": "User Execution (T1204)"},
                    {"tactic": "Discovery", "technique": "System Information Discovery (T1082)"}
                ]
            else:
                attack_stages = [
                    {"tactic": "Reconnaissance", "technique": "Active Scanning (T1595)"},
                    {"tactic": "Discovery", "technique": "System Information Discovery (T1082)"}
                ]
        
        # Create attack flow diagram
        self.create_attack_flow_diagram(finding, attack_stages)
        
        # Add textual description of attack stages
        self.document.add_paragraph("Attack Stages:").bold = True
        for stage in attack_stages:
            stage_para = self.document.add_paragraph(style='List Bullet')
            stage_para.add_run(f"{stage['tactic']}: {stage['technique']}").bold = True
        
        # Add tools used by attackers
        self.document.add_paragraph("Tools Used:").bold = True
        tools_para = self.document.add_paragraph(style='List Bullet')
        
        # Map tools based on finding type
        if "unauthenticated access" in finding["title"].lower() or "signalr" in finding["title"].lower():
            tools_para.add_run("WebSocket Interception Tools (Burp Suite, OWASP ZAP)")
        elif "jquery" in finding["title"].lower() or "javascript" in finding["title"].lower():
            tools_para.add_run("JavaScript Exploitation Frameworks (BeEF, XSS payloads)")
        elif "information disclosure" in finding["title"].lower() or "server information" in finding["title"].lower():
            tools_para.add_run("Reconnaissance Tools (Nmap, Nikto, WhatWeb)")
        else:
            tools_para.add_run("Various penetration testing tools")
        
        # Add references to MITRE ATT&CK
        self.document.add_paragraph("References:").bold = True
        ref_para = self.document.add_paragraph(style='List Bullet')
        ref_para.add_run("MITRE ATT&CK Enterprise Framework: https://attack.mitre.org/matrices/enterprise/")

    def create_attack_flow_diagram(self, finding, attack_stages):
        """Create an attack flow diagram using graphviz"""
        try:
            # Create reports directory if it doesn't exist
            reports_dir = "reports"
            attack_flows_dir = os.path.join(reports_dir, "attack_flows")
            os.makedirs(attack_flows_dir, exist_ok=True)
            
            # Create a unique filename based on finding title
            safe_title = "".join(c if c.isalnum() else "_" for c in finding["title"])
            filename = f"{safe_title}_attack_flow"
            filepath = os.path.join(attack_flows_dir, filename)
            
            # Create the graph
            dot = graphviz.Digraph(comment=f'Attack Flow for {finding["title"]}')
            dot.attr('node', shape='box', style='filled', fillcolor='lightblue')
            
            # Add nodes and edges for each attack stage
            prev_node = None
            for i, stage in enumerate(attack_stages):
                node_id = f"stage_{i}"
                label = f"{stage['tactic']}\n{stage['technique']}"
                dot.node(node_id, label)
                
                if prev_node:
                    dot.edge(prev_node, node_id)
                prev_node = node_id
            
            # Render the graph
            dot.render(filepath, format='png', cleanup=True)
            
            # Add the image to the document
            self.document.add_paragraph("Attack Flow Diagram:").bold = True
            self.document.add_picture(f"{filepath}.png", width=Inches(6))
            
        except Exception as e:
            # If diagram creation fails, add a note
            error_para = self.document.add_paragraph()
            error_para.add_run(f"Attack flow diagram could not be generated: {str(e)}").italic = True

    def add_developer_recommendations(self, finding):
        """Add developer recommendations for the finding"""
        self.add_heading("Developer Recommendations", 3)
        
        # Add implementation guide based on finding type
        if "unauthenticated access" in finding["title"].lower() or "signalr" in finding["title"].lower():
            self.add_heading("SignalR Security Implementation Guide", 4)
            
            # Add authentication implementation
            self.document.add_paragraph("Authentication Implementation:").bold = True
            self.document.add_paragraph("Implement JWT authentication for SignalR hubs:")
            
            # Add code example for server-side configuration
            code_para = self.document.add_paragraph(style='CodeBlock')
            code_para.add_run("// Server-side configuration in Startup.cs").bold = True
            code_para.add_run("""
services.AddAuthentication(options => {
    options.DefaultAuthenticateScheme = JwtBearerDefaults.AuthenticationScheme;
    options.DefaultChallengeScheme = JwtBearerDefaults.AuthenticationScheme;
}).AddJwtBearer(options => {
    options.TokenValidationParameters = new TokenValidationParameters {
        ValidateIssuer = true,
        ValidateAudience = true,
        ValidateLifetime = true,
        ValidateIssuerSigningKey = true,
        ValidIssuer = Configuration["Jwt:Issuer"],
        ValidAudience = Configuration["Jwt:Audience"],
        IssuerSigningKey = new SymmetricSecurityKey(
            Encoding.UTF8.GetBytes(Configuration["Jwt:Key"]))
    };
    
    // Configure the JWT Bearer Events for SignalR
    options.Events = new JwtBearerEvents {
        OnMessageReceived = context => {
            var accessToken = context.Request.Query["access_token"];
            var path = context.HttpContext.Request.Path;
            
            if (!string.IsNullOrEmpty(accessToken) && 
                path.StartsWithSegments("/hub")) {
                context.Token = accessToken;
            }
            return Task.CompletedTask;
        }
    };
});

// Apply authentication to hub
app.UseAuthentication();
app.UseAuthorization();""")
            
            # Add code example for hub implementation
            code_para = self.document.add_paragraph(style='CodeBlock')
            code_para.add_run("// Hub implementation with authorization").bold = True
            code_para.add_run("""
[Authorize]
public class SecureHub : Hub
{
    public async Task SendMessage(string user, string message)
    {
        // Get user identity from JWT claims
        var userId = Context.User.FindFirst(ClaimTypes.NameIdentifier)?.Value;
        
        // Only proceed if user is authenticated
        if (userId != null)
        {
            await Clients.All.SendAsync("ReceiveMessage", user, message);
        }
    }
}""")
            
            # Add implementation steps
            self.document.add_paragraph("Implementation Steps:").bold = True
            steps = [
                "Configure JWT authentication in your application",
                "Apply the [Authorize] attribute to your Hub class",
                "Verify user identity using Context.User in hub methods",
                "Configure client-side to send the JWT token with connection",
                "Test authentication by attempting to connect without a valid token"
            ]
            for step in steps:
                step_para = self.document.add_paragraph(style='List Bullet')
                step_para.add_run(step)
                
        elif "jquery" in finding["title"].lower() or "javascript" in finding["title"].lower():
            self.add_heading("jQuery Security Implementation Guide", 4)
            
            # Add update recommendation
            self.document.add_paragraph("Update jQuery to Latest Version:").bold = True
            self.document.add_paragraph("Replace the current jQuery reference with the latest stable version:")
            
            # Add code example
            code_para = self.document.add_paragraph(style='CodeBlock')
            code_para.add_run("<!-- Replace this -->").bold = True
            code_para.add_run("""
<script src="/js/jquery-1.12.4.min.js"></script>

<!-- With this -->
<script src="https://code.jquery.com/jquery-3.7.1.min.js" 
        integrity="sha256-/JqT3SQfawRcv/BIHPThkBvs0OEvtFFmqPF/lYI/Cxo=" 
        crossorigin="anonymous"></script>""")
            
            # Add dependency management recommendation
            self.document.add_paragraph("Implement Dependency Management:").bold = True
            code_para = self.document.add_paragraph(style='CodeBlock')
            code_para.add_run("// Using npm for dependency management").bold = True
            code_para.add_run("""
// package.json
{
  "dependencies": {
    "jquery": "^3.7.1"
  }
}

// Import in your JavaScript
import $ from 'jquery';""")
            
            # Add implementation steps
            self.document.add_paragraph("Implementation Steps:").bold = True
            steps = [
                "Audit all jQuery usage in your application",
                "Test application functionality with the latest jQuery version",
                "Update all direct script references",
                "Consider implementing a dependency management system",
                "Set up automated security scanning for JavaScript dependencies"
            ]
            for step in steps:
                step_para = self.document.add_paragraph(style='List Bullet')
                step_para.add_run(step)
                
        elif "information disclosure" in finding["title"].lower() or "server information" in finding["title"].lower():
            self.add_heading("Server Header Security Guide", 4)
            
            # Add IIS server configuration
            self.document.add_paragraph("IIS Server Configuration:").bold = True
            self.document.add_paragraph("Remove server information from HTTP headers using web.config:")
            
            # Add code example
            code_para = self.document.add_paragraph(style='CodeBlock')
            code_para.add_run("<!-- web.config for IIS -->").bold = True
            code_para.add_run("""
<system.webServer>
  <security>
    <requestFiltering removeServerHeader="true" />
  </security>
  <httpProtocol>
    <customHeaders>
      <remove name="X-Powered-By" />
      <remove name="Server" />
    </customHeaders>
  </httpProtocol>
</system.webServer>""")
            
            # Add URL Rewrite rule
            self.document.add_paragraph("Using URL Rewrite Module:").bold = True
            code_para = self.document.add_paragraph(style='CodeBlock')
            code_para.add_run("<!-- URL Rewrite rule -->").bold = True
            code_para.add_run("""
<system.webServer>
  <rewrite>
    <outboundRules>
      <rule name="Remove Server header">
        <match serverVariable="RESPONSE_SERVER" pattern=".+" />
        <action type="Rewrite" value="" />
      </rule>
      <rule name="Remove X-Powered-By header">
        <match serverVariable="RESPONSE_X-Powered-By" pattern=".+" />
        <action type="Rewrite" value="" />
      </rule>
    </outboundRules>
  </rewrite>
</system.webServer>""")
            
            # Add implementation steps
            self.document.add_paragraph("Implementation Steps:").bold = True
            steps = [
                "Install URL Rewrite module if using IIS",
                "Add the configuration to your web.config file",
                "Test headers using curl or browser developer tools",
                "Verify all custom headers are properly sanitized",
                "Implement similar configurations for other web servers (Apache, Nginx)"
            ]
            for step in steps:
                step_para = self.document.add_paragraph(style='List Bullet')
                step_para.add_run(step)
        else:
            # Generic security recommendations
            self.add_heading("Security Implementation Guide", 4)
            
            self.document.add_paragraph("General Security Recommendations:").bold = True
            recommendations = [
                "Implement proper input validation for all user inputs",
                "Apply the principle of least privilege for all system components",
                "Use parameterized queries to prevent SQL injection",
                "Implement Content Security Policy (CSP) headers",
                "Enable HTTPS across the entire application",
                "Implement proper session management and timeout policies",
                "Regularly update all dependencies and frameworks"
            ]
            for rec in recommendations:
                rec_para = self.document.add_paragraph(style='List Bullet')
                rec_para.add_run(rec)
        
        # Add page break after developer recommendations
        self.document.add_page_break()

    def add_heading(self, text, level):
        return self.document.add_heading(text, level=level)

    def generate(self):
        # Create reports directory if it doesn't exist
        reports_dir = "reports"
        os.makedirs(reports_dir, exist_ok=True)
        
        # Generate the report
        self.add_title_page()
        self.add_executive_summary()
        self.add_findings()
        self.add_developer_recommendations()
        
        # Save the document
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"VAPT_Report_{timestamp}.docx"
        filepath = os.path.join(reports_dir, filename)
        self.document.save(filepath)
        
        return filepath

def generate_report(report_data):
    generator = VAPTReportGenerator(report_data)
    return generator.generate()